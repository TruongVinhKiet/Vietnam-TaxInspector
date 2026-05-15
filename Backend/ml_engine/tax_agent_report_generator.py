import io
import json
import logging
from datetime import datetime
from typing import Any, Dict, List

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib import colors
from reportlab.lib.units import inch as rl_inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger(__name__)

# ── Register Unicode font for Vietnamese PDF support ──
_VIET_FONT = "Helvetica"  # fallback
_VIET_FONT_BOLD = "Helvetica-Bold"
try:
    import os
    _font_candidates = [
        os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "arial.ttf"),
        os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "arialbd.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]
    _regular = next((f for f in _font_candidates if os.path.isfile(f) and "bold" not in f.lower() and "bd" not in f.lower()), None)
    _bold = next((f for f in _font_candidates if os.path.isfile(f) and ("bold" in f.lower() or "bd" in f.lower())), None)
    if _regular:
        pdfmetrics.registerFont(TTFont("VietFont", _regular))
        _VIET_FONT = "VietFont"
        if _bold:
            pdfmetrics.registerFont(TTFont("VietFontBold", _bold))
            _VIET_FONT_BOLD = "VietFontBold"
        else:
            _VIET_FONT_BOLD = "VietFont"
        logger.info("Registered Vietnamese PDF font: %s", _regular)
except Exception as e:
    logger.warning("Could not register Vietnamese font, PDF may have missing chars: %s", e)

class TaxReportGenerator:
    """Generates official tax inspection reports from session context."""

    def __init__(self, context_data: Dict[str, Any]):
        self.context = context_data
        self.batch_data = context_data.get("batch_data") or {}
        self.vat_snapshot = context_data.get("vat_snapshot") or {}
        self.facts = context_data.get("facts") or []
        self.recent_turns = context_data.get("recent_turns") or []
        self.summary = context_data.get("session_summary") or "Chưa có tóm tắt phiên."

    def generate_docx(self) -> bytes:
        doc = Document()
        
        # Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập – Tự do – Hạnh phúc")
        run.bold = True
        p.add_run("\n─────────────────────────")
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BÁO CÁO KẾT QUẢ PHÂN TÍCH RỦI RO THUẾ\n(Trích xuất từ Hệ thống TaxInspector AI)")
        run.bold = True
        run.font.size = Pt(14)
        
        p = doc.add_paragraph(f"Số: {datetime.now().strftime('%Y%m%d%H%M')}/BC-TTAI")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph(f"Ngày: {datetime.now().strftime('%d/%m/%Y')}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # I. Thông tin chung
        doc.add_heading("I. THÔNG TIN CHUNG", level=1)
        doc.add_paragraph(f"Tóm tắt ngữ cảnh phân tích: {self.summary}")

        has_data = False

        # II. Fraud / Lô rủi ro
        if self.batch_data and self.batch_data.get("companies"):
            has_data = True
            doc.add_heading("II. KẾT QUẢ PHÂN TÍCH GIAN LẬN", level=1)
            companies = self.batch_data.get("companies", [])
            doc.add_paragraph(f"Tổng số doanh nghiệp phân tích: {len(companies)}")
            
            # Bảng top 5
            top_5 = self.batch_data.get("top_risky", companies[:5])
            if top_5:
                doc.add_heading("Top doanh nghiệp rủi ro cao nhất:", level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = "MST"
                hdr[1].text = "Tên doanh nghiệp"
                hdr[2].text = "Mức rủi ro"
                hdr[3].text = "Điểm"
                for c in top_5:
                    row = table.add_row().cells
                    row[0].text = str(c.get("tax_code", ""))
                    row[1].text = str(c.get("company_name", ""))
                    row[2].text = str(c.get("risk_level", "Unknown"))
                    row[3].text = str(c.get("risk_score", 0))
            
            # Chart 1: Risk Distribution
            try:
                by_level = self.batch_data.get("by_level", {})
                viz_fraud = self.context.get("viz_data", {}).get("fraud", {})
                if not by_level and viz_fraud.get("risk_distribution"):
                    by_level = viz_fraud["risk_distribution"]
                    
                if by_level:
                    fig, ax = plt.subplots(figsize=(5, 3))
                    labels = list(by_level.keys())
                    sizes = list(by_level.values())
                    ax.bar(labels, sizes, color=['#ef4444', '#f97316', '#eab308', '#22c55e'][:len(labels)])
                    ax.set_title("Phân bố mức độ rủi ro")
                    ax.set_ylabel("Số lượng DN")
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(4.5))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting fraud risk dist: {e}")

            # Chart 2: Yearly Trend (from viz_data)
            try:
                viz_fraud = self.context.get("viz_data", {}).get("fraud", {})
                trend = viz_fraud.get("yearly_trend", [])
                if trend:
                    fig, ax = plt.subplots(figsize=(5, 2.5))
                    labels = [str(x.get("year", "")) for x in trend]
                    values = [x.get("avg_risk", 0) for x in trend]
                    ax.plot(labels, values, marker='o', linestyle='-', color='#ec4899')
                    ax.set_title("Xu hướng rủi ro theo năm")
                    ax.set_ylabel("Điểm rủi ro TB")
                    ax.grid(True, linestyle='--', alpha=0.5)
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(4.5))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting fraud trend: {e}")

            # Chart 3: Radar — Hồ sơ rủi ro đa chiều (fraud risk factors)
            try:
                viz_fraud = self.context.get("viz_data", {}).get("fraud", {})
                radar = viz_fraud.get("radar", {})
                radar_labels = radar.get("labels", [])
                radar_values = radar.get("values", [])
                if radar_labels and radar_values:
                    import numpy as np
                    angles = np.linspace(0, 2 * np.pi, len(radar_labels), endpoint=False).tolist()
                    radar_values_c = radar_values + [radar_values[0]]
                    angles += [angles[0]]
                    fig, ax = plt.subplots(figsize=(4, 4), subplot_kw=dict(polar=True))
                    ax.fill(angles, radar_values_c, color='#0d9488', alpha=0.25)
                    ax.plot(angles, radar_values_c, color='#0d9488', linewidth=2)
                    ax.set_xticks(angles[:-1])
                    ax.set_xticklabels(radar_labels, fontsize=7)
                    ax.set_title("Hồ sơ rủi ro đa chiều", pad=15)
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(3.8))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting radar: {e}")

            # Chart 4: Donut — Cơ cấu rủi ro
            try:
                by_level = self.batch_data.get("by_level", {})
                if not by_level:
                    viz_fraud = self.context.get("viz_data", {}).get("fraud", {})
                    by_level = viz_fraud.get("risk_distribution", {})
                if by_level:
                    fig, ax = plt.subplots(figsize=(4, 3))
                    labels = list(by_level.keys())
                    sizes = list(by_level.values())
                    colors_pie = ['#ef4444', '#f97316', '#eab308', '#22c55e'][:len(labels)]
                    wedges, texts, autotexts = ax.pie(
                        sizes, labels=labels, colors=colors_pie,
                        autopct='%1.1f%%', startangle=90,
                        wedgeprops=dict(width=0.45, edgecolor='white')
                    )
                    ax.set_title("Cơ cấu phân bố rủi ro")
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(4))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting donut: {e}")

            # Chart 5: Scatter — Doanh thu so với rủi ro
            try:
                viz_fraud = self.context.get("viz_data", {}).get("fraud", {})
                scatter_pts = viz_fraud.get("revenue_risk_scatter", [])
                if scatter_pts:
                    fig, ax = plt.subplots(figsize=(5, 3))
                    xs = [p.get("revenue", p.get("x", 0)) for p in scatter_pts]
                    ys = [p.get("risk_score", p.get("y", 0)) for p in scatter_pts]
                    ax.scatter(xs, ys, c='#6366f1', alpha=0.6, edgecolors='white', s=40)
                    ax.set_xlabel("Doanh thu (VND)")
                    ax.set_ylabel("Điểm rủi ro")
                    ax.set_title("Tương quan Doanh thu — Rủi ro")
                    ax.grid(True, linestyle='--', alpha=0.4)
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(4.5))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting scatter: {e}")

        # III. VAT Network
        if self.vat_snapshot and self.vat_snapshot.get("edges"):
            has_data = True
            doc.add_heading("III. KẾT QUẢ TRUY VẾT MẠNG LƯỚI VAT", level=1)
            summary = self.vat_snapshot.get("summary", {})
            doc.add_paragraph(f"Phân tích {summary.get('invoices', 0)} hóa đơn giữa {summary.get('companies', 0)} doanh nghiệp.")
            
            edges = self.vat_snapshot.get("top_invoice_risks") or self.vat_snapshot.get("edges") or []
            if edges:
                doc.add_heading("Các giao dịch/cạnh rủi ro cao:", level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = "Bên bán"
                hdr[1].text = "Bên mua"
                hdr[2].text = "Giá trị"
                hdr[3].text = "Độ rủi ro"
                for e in edges[:10]:
                    row = table.add_row().cells
                    row[0].text = str(e.get("seller_tax_code") or e.get("source") or "")
                    row[1].text = str(e.get("buyer_tax_code") or e.get("target") or "")
                    row[2].text = str(e.get("amount") or e.get("value") or 0)
                    row[3].text = str(e.get("edge_risk_score") or e.get("risk_score") or "")

        # IV. Macro Simulation
        viz_data = self.context.get("viz_data") or {}
        macro = self.context.get("macro_forecast") or viz_data.get("macro") or {}
        if macro:
            has_data = True
            doc.add_heading("IV. KẾT QUẢ MÔ PHỎNG VĨ MÔ", level=1)
            summary_macro = str(macro.get("scenario_name", "Kịch bản tùy chỉnh"))
            doc.add_paragraph(f"Kịch bản: {summary_macro}")
            
            kpis = macro.get("kpis", {})
            if kpis:
                doc.add_paragraph(f"Dự kiến doanh nghiệp rủi ro cao: {kpis.get('simulated_high_risk_count', 0)}")
                doc.add_paragraph(f"Tỷ lệ nợ đọng dự báo: {kpis.get('simulated_delinquency_rate', 0)}%")
                doc.add_paragraph(f"Doanh thu mô phỏng: {kpis.get('simulated_total_revenue', 0):,} VND")
                
            # Chart: Quarterly Projection
            try:
                proj = macro.get("quarterly_projection", [])
                if proj:
                    fig, ax = plt.subplots(figsize=(5.5, 3))
                    labels = [str(x.get("quarter", x.get("label", ""))) for x in proj]
                    values = [float(x.get("simulated_value", x.get("value", 0))) for x in proj]
                    ax.plot(labels, values, marker='s', color='#8b5cf6')
                    ax.set_title("Mô phỏng dự thu ngân sách")
                    ax.set_ylabel("Giá trị (VND)")
                    ax.grid(True, linestyle='--', alpha=0.5)
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(5))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting macro chart: {e}")

            # Industry Impact Table (from simulation page)
            industries = macro.get("industry_impacts", [])
            if industries:
                doc.add_heading("Ma trận tác động theo ngành:", level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = "Ngành"
                hdr[1].text = "NĐ gốc (%)"
                hdr[2].text = "NĐ mô phỏng (%)"
                hdr[3].text = "Delta (%)"
                for ind in industries[:10]:
                    row = table.add_row().cells
                    row[0].text = str(ind.get("industry", ""))
                    row[1].text = f"{ind.get('baseline_delinquency_rate', 0):.2f}"
                    row[2].text = f"{ind.get('simulated_delinquency_rate', 0):.2f}"
                    row[3].text = f"{ind.get('delta_pct', 0):.2f}"
        
        # V. Delinquency Timeline
        delinq = viz_data.get("delinquency_timeline") or {}
        if delinq:
            has_data = True
            doc.add_heading("V. DỰ BÁO NỢ ĐỌNG (DELINQUENCY)", level=1)
            doc.add_paragraph(f"Mô hình sử dụng: {delinq.get('dl_architecture', 'Chuẩn')}")
            
            # Chart
            try:
                ml_values = delinq.get("ml_values", [])
                dl_values = delinq.get("dl_values", [])
                labels = delinq.get("labels", ["30 ngày", "60 ngày", "90 ngày"])
                
                if ml_values or dl_values:
                    fig, ax = plt.subplots(figsize=(5.5, 3))
                    import numpy as np
                    x = np.arange(len(labels))
                    width = 0.35
                    if ml_values:
                        ax.bar(x - width/2, ml_values, width, label='ML Pipeline', color='#94a3b8')
                    if dl_values:
                        ax.bar(x + width/2, dl_values, width, label='Deep Learning', color='#3b82f6')
                    ax.set_ylabel('Xác suất nợ đọng (%)')
                    ax.set_title('So sánh dự báo nợ đọng')
                    ax.set_xticks(x)
                    ax.set_xticklabels(labels)
                    ax.legend()
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(5))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting delinquency chart: {e}")

        # VI. Facts & Legal Context
        if self.facts:
            has_data = True
            doc.add_heading("VI. TỔNG HỢP KIẾN NGHỊ VÀ TƯ VẤN (TỪ HỆ THỐNG)", level=1)
            for f in self.facts:
                mode = f.get("mode", "Chung")
                claim = f.get("claim_text", "")
                if claim:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"[{mode.upper()}] ").bold = True
                    p.add_run(claim)
                    
            legal_facts = [f for f in self.facts if f.get("mode") == "legal" or f.get("intent") == "legal"]
            if legal_facts:
                doc.add_heading("Tham chiếu pháp lý chi tiết:", level=2)
                for f in legal_facts:
                    val = f.get("value_json", {})
                    if val and isinstance(val, dict):
                        doc.add_paragraph(f"- {val.get('title', 'Tài liệu')}: {val.get('snippet', '')}", style='Body Text')
        if not has_data:
            doc.add_paragraph("Chưa có đủ dữ liệu phân tích chuyên sâu trong phiên làm việc này.")

        # Footer
        doc.add_paragraph("\n")
        p = doc.add_paragraph("KẾT LUẬN VÀ KIẾN NGHỊ", style='Heading 2')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Cơ quan thuế cần xem xét các điểm rủi ro nêu trên để có biện pháp nghiệp vụ phù hợp (thanh tra, kiểm tra, hoặc cảnh báo).")

        p = doc.add_paragraph("\n\nNgười lập báo cáo\nHệ thống TaxInspector AI")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        out_stream = io.BytesIO()
        doc.save(out_stream)
        return out_stream.getvalue()

    def generate_pdf(self) -> bytes:
        out_stream = io.BytesIO()
        doc = SimpleDocTemplate(out_stream, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = []

        # Override all styles to use Vietnamese-capable font
        for style_name in styles.byName:
            styles[style_name].fontName = _VIET_FONT
        styles['Heading1'].fontName = _VIET_FONT_BOLD
        styles['Heading2'].fontName = _VIET_FONT_BOLD
        styles['Heading3'].fontName = _VIET_FONT_BOLD

        title_style = ParagraphStyle(
            'VietTitle',
            parent=styles['Heading1'],
            fontName=_VIET_FONT_BOLD,
            alignment=1, # Center
            spaceAfter=20
        )
        
        elements.append(Paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM<br/>Độc lập – Tự do – Hạnh phúc", title_style))
        elements.append(Paragraph("BÁO CÁO KẾT QUẢ PHÂN TÍCH RỦI RO THUẾ", title_style))
        
        elements.append(Paragraph(f"Tóm tắt: {self.summary}", styles['Normal']))
        elements.append(Spacer(1, 20))

        if self.batch_data and self.batch_data.get("top_risky"):
            elements.append(Paragraph("I. GIAN LẬN", styles['Heading2']))
            data = [["MST", "Tên DN", "Rủi ro"]]
            for c in self.batch_data.get("top_risky", [])[:5]:
                data.append([str(c.get("tax_code")), str(c.get("company_name"))[:20], str(c.get("risk_level"))])
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
            ]))
            elements.append(t)
            
            # Trend Chart
            viz_fraud = self.context.get("viz_data", {}).get("fraud", {})
            trend = viz_fraud.get("yearly_trend", [])
            if trend:
                elements.append(Spacer(1, 10))
                try:
                    fig, ax = plt.subplots(figsize=(5, 2.5))
                    labels = [str(x.get("year", "")) for x in trend]
                    values = [x.get("avg_risk", 0) for x in trend]
                    ax.plot(labels, values, marker='o', linestyle='-', color='#ec4899')
                    ax.set_title("Xu hướng rủi ro theo năm")
                    ax.set_ylabel("Điểm rủi ro TB")
                    ax.grid(True, linestyle='--', alpha=0.5)
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    elements.append(RLImage(img_stream, width=5*rl_inch, height=2.5*rl_inch))
                    plt.close(fig)
                except Exception as e:
                    logger.error(f"Error plotting fraud trend pdf: {e}")

        # II. VAT Network
        edges = self.vat_snapshot.get("top_invoice_risks") or self.vat_snapshot.get("edges") or []
        if edges:
            elements.append(Spacer(1, 20))
            elements.append(Paragraph("II. TRUY VẾT MẠNG LƯỚI VAT", styles['Heading2']))
            data = [["Bên bán", "Bên mua", "Giá trị", "Độ rủi ro"]]
            for e in edges[:8]:
                data.append([
                    str(e.get("seller_tax_code") or e.get("source") or ""),
                    str(e.get("buyer_tax_code") or e.get("target") or ""),
                    str(e.get("amount") or e.get("value") or 0),
                    str(e.get("edge_risk_score") or e.get("risk_score") or "")
                ])
            t = Table(data)
            t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 1, colors.black)]))
            elements.append(t)

        # III. Macro
        macro = self.context.get("macro_forecast") or self.context.get("viz_data", {}).get("macro") or {}
        if macro:
            elements.append(Spacer(1, 20))
            elements.append(Paragraph("III. KẾT QUẢ MÔ PHỎNG VĨ MÔ", styles['Heading2']))
            elements.append(Paragraph(f"Kịch bản: {macro.get('scenario_name', 'Tùy chỉnh')}", styles['Normal']))
            kpis = macro.get("kpis", {})
            if kpis:
                data = [["Chỉ số", "Giá trị"]]
                data.append(["Doanh nghiệp rủi ro cao", str(kpis.get('simulated_high_risk_count', 0))])
                data.append(["Tỷ lệ nợ đọng", f"{kpis.get('simulated_delinquency_rate', 0)}%"])
                data.append(["Doanh thu mô phỏng", f"{kpis.get('simulated_total_revenue', 0):,} VND"])
                t = Table(data)
                t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 1, colors.black)]))
                elements.append(t)
                
            # Quarterly projection chart
            proj = macro.get("quarterly_projection", [])
            if proj:
                elements.append(Spacer(1, 10))
                try:
                    fig, ax = plt.subplots(figsize=(5, 2.5))
                    labels = [str(x.get("quarter", x.get("label", ""))) for x in proj]
                    values = [float(x.get("simulated_value", x.get("value", 0))) for x in proj]
                    ax.plot(labels, values, marker='s', color='#8b5cf6')
                    ax.set_title("Mô phỏng dự thu ngân sách")
                    ax.set_ylabel("Giá trị (VND)")
                    ax.grid(True, linestyle='--', alpha=0.5)
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    elements.append(RLImage(img_stream, width=5*rl_inch, height=2.5*rl_inch))
                    plt.close(fig)
                except Exception as e:
                    logger.error(f"Error plotting macro pdf chart: {e}")

        # IV. Delinquency Timeline
        delinq = self.context.get("viz_data", {}).get("delinquency_timeline") or {}
        if delinq:
            elements.append(Spacer(1, 20))
            elements.append(Paragraph("IV. DỰ BÁO NỢ ĐỌNG", styles['Heading2']))
            elements.append(Paragraph(f"Mô hình sử dụng: {delinq.get('dl_architecture', 'Chuẩn')}", styles['Normal']))
            
            try:
                ml_values = delinq.get("ml_values", [])
                dl_values = delinq.get("dl_values", [])
                labels = delinq.get("labels", ["30 ngày", "60 ngày", "90 ngày"])
                
                if ml_values or dl_values:
                    fig, ax = plt.subplots(figsize=(5, 2.5))
                    import numpy as np
                    x = np.arange(len(labels))
                    width = 0.35
                    if ml_values: ax.bar(x - width/2, ml_values, width, label='ML Pipeline')
                    if dl_values: ax.bar(x + width/2, dl_values, width, label='Deep Learning')
                    ax.set_ylabel('Xác suất (%)')
                    ax.set_xticks(x)
                    ax.set_xticklabels(labels)
                    ax.legend()
                    img_stream = io.BytesIO()
                    plt.savefig(img_stream, format='png', bbox_inches='tight')
                    img_stream.seek(0)
                    elements.append(RLImage(img_stream, width=5*rl_inch, height=2.5*rl_inch))
                    plt.close(fig)
            except Exception as e:
                logger.error(f"Error plotting delinquency pdf chart: {e}")

        # V. Facts & Legal
        if self.facts:
            elements.append(Spacer(1, 20))
            elements.append(Paragraph("V. TƯ VẤN & KIẾN NGHỊ", styles['Heading2']))
            for f in self.facts:
                if f.get("claim_text"):
                    elements.append(Paragraph(f"- [{f.get('mode', 'Chung')}] {f.get('claim_text')}", styles['Normal']))

        doc.build(elements)
        return out_stream.getvalue()
